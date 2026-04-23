"""
RAG 知识库助手 - 文档加载器模块

支持多种文档格式的解析：
- PDF (使用 PyMuPDF)
- Markdown (使用正则解析)
- TXT (纯文本)
- DOCX (使用 python-docx)

统一接口：load_document(file_path) -> list[Document]
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Union


@dataclass
class Document:
    """文档内容数据类
    
    Attributes:
        page_content: 文本内容
        metadata: 元数据字典，包含 source、page 等信息
    """
    page_content: str
    metadata: dict = field(default_factory=dict)
    
    def __post_init__(self):
        """初始化后处理，确保 metadata 是字典"""
        if self.metadata is None:
            self.metadata = {}


class DocumentLoaderError(Exception):
    """文档加载器异常基类"""
    pass


class UnsupportedFileTypeError(DocumentLoaderError):
    """不支持的文件类型异常"""
    pass


class FileNotFoundError(DocumentLoaderError):
    """文件不存在异常"""
    pass


class DocumentParseError(DocumentLoaderError):
    """文档解析异常"""
    pass


class DocumentLoader:
    """文档加载器类
    
    支持 PDF、Markdown、TXT、DOCX 格式的文档解析。
    
    Example:
        >>> loader = DocumentLoader()
        >>> documents = loader.load("document.pdf")
        >>> for doc in documents:
        ...     print(f"Page {doc.metadata.get('page')}: {doc.page_content[:100]}")
    """
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}
    
    def __init__(self):
        """初始化文档加载器"""
        self._parsers = {
            ".pdf": self._parse_pdf,
            ".md": self._parse_markdown,
            ".txt": self._parse_txt,
            ".docx": self._parse_docx,
        }
    
    def load(self, file_path: Union[str, Path]) -> list[Document]:
        """加载并解析文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document 对象列表
            
        Raises:
            FileNotFoundError: 文件不存在
            UnsupportedFileTypeError: 不支持的文件类型
            DocumentParseError: 文档解析失败
        """
        file_path = Path(file_path)
        
        # 检查文件是否存在
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件类型
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型: {ext}。"
                f"支持的类型: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # 调用对应的解析器
        parser = self._parsers[ext]
        try:
            return parser(file_path)
        except Exception as e:
            raise DocumentParseError(f"解析文件失败 {file_path}: {str(e)}") from e
    
    def _parse_pdf(self, file_path: Path) -> list[Document]:
        """解析 PDF 文件
        
        使用 PyMuPDF (fitz) 提取文本，保留页码信息。
        过滤掉内容过短的页面（如目录页、空白页）。
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            每页一个 Document 对象
        """
        import fitz  # PyMuPDF
        
        documents = []
        
        with fitz.open(file_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text = page.get_text()
                
                # 清理文本（移除多余空白）
                text = self._clean_text(text)
                
                # 过滤条件：内容非空且长度足够（至少 50 个字符）
                # 避免目录页、页眉页脚等短内容干扰检索
                if text.strip() and len(text.strip()) >= 50:
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "page": page_num + 1,  # 页码从 1 开始
                            "file_type": "pdf",
                        }
                    )
                    documents.append(doc)
        
        return documents
    
    def _parse_markdown(self, file_path: Path) -> list[Document]:
        """解析 Markdown 文件
        
        按标题层级分割文档，保留标题结构信息。
        
        Args:
            file_path: Markdown 文件路径
            
        Returns:
            按标题分割的 Document 对象列表
        """
        content = file_path.read_text(encoding="utf-8")
        
        # 按标题分割（匹配 # ## ### 等）
        # 使用正则表达式匹配标题行
        heading_pattern = r'^(#{1,6}\s+.+)$'
        
        documents = []
        current_section = []
        current_heading = ""
        section_index = 0
        
        lines = content.split('\n')
        
        for line in lines:
            if re.match(heading_pattern, line.strip()):
                # 保存之前的段落
                if current_section:
                    text = '\n'.join(current_section).strip()
                    if text:
                        doc = Document(
                            page_content=text,
                            metadata={
                                "source": file_path.name,
                                "section": section_index,
                                "heading": current_heading,
                                "file_type": "markdown",
                            }
                        )
                        documents.append(doc)
                    section_index += 1
                
                current_heading = line.strip()
                current_section = [line]
            else:
                current_section.append(line)
        
        # 处理最后一个段落
        if current_section:
            text = '\n'.join(current_section).strip()
            if text:
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "section": section_index,
                        "heading": current_heading,
                        "file_type": "markdown",
                    }
                )
                documents.append(doc)
        
        # 如果没有按标题分割（没有标题），返回整个文档
        if not documents:
            text = self._clean_text(content)
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "section": 0,
                        "file_type": "markdown",
                    }
                ))
        
        return documents
    
    def _parse_txt(self, file_path: Path) -> list[Document]:
        """解析 TXT 文件
        
        按段落分割文本。
        
        Args:
            file_path: TXT 文件路径
            
        Returns:
            按段落分割的 Document 对象列表
        """
        content = file_path.read_text(encoding="utf-8")
        
        # 按段落分割（空行分隔）
        paragraphs = re.split(r'\n\s*\n', content)
        
        documents = []
        for idx, para in enumerate(paragraphs):
            text = self._clean_text(para)
            if text.strip():  # 只添加非空段落
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "paragraph": idx + 1,
                        "file_type": "txt",
                    }
                )
                documents.append(doc)
        
        # 如果没有有效段落，返回整个文档
        if not documents:
            text = self._clean_text(content)
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "paragraph": 1,
                        "file_type": "txt",
                    }
                ))
        
        return documents
    
    def _parse_docx(self, file_path: Path) -> list[Document]:
        """解析 DOCX 文件
        
        使用 python-docx 提取文本，按段落分割。
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            按段落分割的 Document 对象列表
        """
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        documents = []
        para_index = 0
        
        for para in doc.paragraphs:
            text = self._clean_text(para.text)
            if text.strip():  # 只添加非空段落
                para_index += 1
                document = Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "paragraph": para_index,
                        "file_type": "docx",
                    }
                )
                documents.append(document)
        
        # 如果没有有效段落，尝试提取表格内容
        if not documents:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = self._clean_text(cell.text)
                        if cell_text.strip():
                            row_text.append(cell_text)
                    
                    if row_text:
                        text = " | ".join(row_text)
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source": file_path.name,
                                "table": table_idx + 1,
                                "row": row_idx + 1,
                                "file_type": "docx",
                            }
                        ))
        
        return documents
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """清理文本
        
        移除多余空白、统一换行符等。
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除行尾空白
        lines = [line.rstrip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # 移除多余空行（最多保留一个空行）
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()


def load_document(file_path: Union[str, Path]) -> list[Document]:
    """便捷函数：加载文档
    
    这是 DocumentLoader.load() 的快捷方式。
    
    Args:
        file_path: 文件路径
        
    Returns:
        Document 对象列表
        
    Example:
        >>> docs = load_document("document.pdf")
        >>> for doc in docs:
        ...     print(doc.page_content[:100])
    """
    loader = DocumentLoader()
    return loader.load(file_path)
